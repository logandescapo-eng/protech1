"""Generate the ProTech academic journal-style article (DOCX)."""

import sys
from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

ROOT = Path(__file__).resolve().parent.parent
OUTPUT = ROOT / "documentattion" / "ProTech_Academic_Article_Patrice_Junior.docx"

sys.path.insert(0, str(ROOT / "scripts"))
from diagrams import generate_all_diagrams  # noqa: E402

BLUE_DARK = "1E3A8A"
BLUE_MID = "2563EB"
BLUE_LIGHT = "DBEAFE"
BLUE_PALE = "EFF6FF"
WHITE = "FFFFFF"

AUTHOR = "Manye Patrice"
PROGRAMME = "IT Innovation in Business — Master's Degree"
COURSES = "Web Technology in Business · Web Application Development"
EMAIL = "manye.patrice@university.edu"
REPO = "github.com/logandescapo-eng/protech1"
JOURNAL = "International Journal of Web Engineering & Business Information Systems"
ARTICLE_HEADER = f"ProTech Academic Article  |  {AUTHOR}  |  {PROGRAMME}"
FOOTER_LEFT = f"{AUTHOR}  |  Academic Article"


def add_page_number_field(paragraph):
    run = paragraph.add_run()
    for fld_type in ("begin", "instr", "separate", "end"):
        if fld_type == "instr":
            instr = OxmlElement("w:instrText")
            instr.set(qn("xml:space"), "preserve")
            instr.text = " PAGE "
            run._r.append(instr)
        else:
            fld = OxmlElement("w:fldChar")
            fld.set(qn("w:fldCharType"), fld_type)
            run._r.append(fld)


def setup_headers_footers(doc):
    for section in doc.sections:
        section.different_first_page_header_footer = True
        section.header_distance = Inches(0.45)
        section.footer_distance = Inches(0.4)
        section.first_page_header.paragraphs[0].clear()
        section.first_page_footer.paragraphs[0].clear()

        hp = section.header.paragraphs[0]
        hp.clear()
        hp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        hr = hp.add_run(ARTICLE_HEADER)
        hr.font.size = Pt(8)
        hr.font.color.rgb = RGBColor(0x1E, 0x3A, 0x8A)
        hr.italic = True

        fp = section.footer.paragraphs[0]
        fp.clear()
        fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        fl = fp.add_run(FOOTER_LEFT + "  |  Page ")
        fl.font.size = Pt(8)
        fl.font.color.rgb = RGBColor(0x64, 0x74, 0x8B)
        add_page_number_field(fp)


def set_cell_shading(cell, hex_color):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), hex_color)
    shd.set(qn("w:val"), "clear")
    tc_pr.append(shd)


def set_cell_margins(cell, top=80, bottom=80, left=120, right=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    mar = OxmlElement("w:tcMar")
    for side, val in (("top", top), ("bottom", bottom), ("start", left), ("end", right)):
        el = OxmlElement(f"w:{side}")
        el.set(qn("w:w"), str(val))
        el.set(qn("w:type"), "dxa")
        mar.append(el)
    tc_pr.append(mar)


def write_cell(cell, text, bold=False, size=10, color=None, align=WD_ALIGN_PARAGRAPH.LEFT):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = align
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    run.font.name = "Times New Roman"
    if color:
        run.font.color.rgb = RGBColor.from_string(color) if isinstance(color, str) and len(color) == 6 else color


def body(doc, text, size=11, space_after=8, italic=False):
    text = " ".join(text.split())
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.line_spacing = 1.25
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.font.name = "Times New Roman"
    run.italic = italic
    return p


def section_heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0x1E, 0x3A, 0x8A)
        run.font.name = "Times New Roman"
    return h


def caption(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.font.size = Pt(9)
    run.font.italic = True
    run.font.color.rgb = RGBColor(0x64, 0x74, 0x8B)
    run.font.name = "Times New Roman"
    p.paragraph_format.space_after = Pt(10)


def styled_table(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    for i, h in enumerate(headers):
        set_cell_shading(table.rows[0].cells[i], BLUE_DARK)
        set_cell_margins(table.rows[0].cells[i])
        write_cell(table.rows[0].cells[i], h, bold=True, size=9,
                   color=RGBColor(0xFF, 0xFF, 0xFF), align=WD_ALIGN_PARAGRAPH.CENTER)
    for ri, row in enumerate(rows):
        cells = table.add_row().cells
        fill = BLUE_PALE if ri % 2 == 0 else WHITE
        for ci, val in enumerate(row):
            set_cell_shading(cells[ci], fill)
            set_cell_margins(cells[ci])
            write_cell(cells[ci], str(val), size=9)
    if col_widths:
        for row in table.rows:
            for i, w in enumerate(col_widths):
                row.cells[i].width = Inches(w)
    tbl_pr = table._tbl.tblPr
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), "4")
        el.set(qn("w:color"), "CBD5E1")
        borders.append(el)
    tbl_pr.append(borders)
    return table


def code_box(doc, text):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.rows[0].cells[0]
    set_cell_shading(cell, "F1F5F9")
    set_cell_margins(cell, top=100, bottom=100, left=140, right=140)
    cell.text = ""
    for i, line in enumerate(text.strip().split("\n")):
        p = cell.paragraphs[0] if i == 0 else cell.add_paragraph()
        run = p.add_run(line)
        run.font.name = "Consolas"
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0x1E, 0x3A, 0x8A)


def add_figure(doc, path, width=5.8):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if Path(path).exists():
        p.add_run().add_picture(str(path), width=Inches(width))


def title_page(doc):
    banner = doc.add_table(rows=1, cols=1)
    banner.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = banner.rows[0].cells[0]
    set_cell_shading(cell, BLUE_DARK)
    set_cell_margins(cell, top=160, bottom=160, left=180, right=180)
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for text, size, bold in [
        (JOURNAL, 10, False),
        (f"Vol. 12  |  Issue 2  |  {date.today().strftime('%B %Y')}  |  ISSN: 2456-XXXX", 9, False),
        ("", 6, False),
        ("ProTech: A Containerised Django Marketplace for Professional Services", 14, True),
        ("with Escrow Payments, RBAC, and Two-Sided Monetisation", 12, True),
        ("", 6, False),
        ("An End-of-Year Master's Project in IT Innovation in Business", 11, False),
    ]:
        r = p.add_run(text + "\n")
        r.font.size = Pt(size)
        r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        r.bold = bold
        r.font.name = "Times New Roman"

    meta = doc.add_table(rows=4, cols=1)
    meta.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, line in enumerate([
        AUTHOR,
        PROGRAMME,
        COURSES,
        f"Email: {EMAIL}  |  Repository: {REPO}",
    ]):
        set_cell_shading(meta.rows[i].cells[0], BLUE_PALE if i % 2 else WHITE)
        write_cell(meta.rows[i].cells[0], line, size=11 if i == 0 else 10,
                   bold=(i == 0), align=WD_ALIGN_PARAGRAPH.CENTER)

    doc.add_page_break()


def abstract_section(doc):
    section_heading(doc, "Abstract", 1)
    abstract_text = (
        "This paper presents ProTech, a full-stack professional services marketplace developed as an "
        "end-of-year master's project within the IT Innovation in Business programme. The platform connects "
        "clients with skilled service workers through a unified web application supporting discovery, booking, "
        "messaging, post-job reviews, and escrow-protected payments. Unlike an earlier PHP/MySQL prototype, "
        "the current system uses Django 5.2.15, PostgreSQL 15, Redis 7, Gunicorn, and Nginx via Docker Compose. "
        "Security is enforced through PBKDF2 hashing, CSRF middleware, ORM parameterisation, template "
        "auto-escaping, and a custom @user_type_required decorator. The escrow subsystem in escrow/services.py "
        "manages wallets, held funds, platform commission (default five percent), and auditable transactions using "
        "Decimal arithmetic and atomic database transactions. The paper covers architecture, business model, "
        "data design, DevOps, testing (seventeen automated tests, GitHub Actions CI), limitations, and future work."
    )
    box = doc.add_table(rows=1, cols=1)
    cell = box.rows[0].cells[0]
    set_cell_shading(cell, "F8FAFC")
    set_cell_margins(cell, top=120, bottom=120, left=160, right=160)
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    run = p.add_run(abstract_text)
    run.font.size = Pt(11)
    run.font.name = "Times New Roman"
    # left accent bar via nested table border trick — simple blue top border on cell
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = OxmlElement("w:tcBorders")
    left = OxmlElement("w:left")
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), "24")
    left.set(qn("w:color"), BLUE_MID)
    borders.append(left)
    tc_pr.append(borders)

    body(doc,
         "Keywords: Django, marketplace platform, escrow payments, role-based access control, "
         "Docker Compose, PostgreSQL, Redis caching, two-sided markets, full-stack web development.",
         size=10, italic=True)
    doc.add_page_break()


def references_section(doc):
    section_heading(doc, "References", 1)
    refs = [
        "Django Software Foundation. (2024). Django Documentation (v5.2). https://docs.djangoproject.com/",
        "Docker Inc. (2024). Docker Compose Specification. https://docs.docker.com/compose/",
        "Edelman, B., Luca, M., & Svirsky, D. (2012). From niches to riches. Harvard Business School Working Paper.",
        "Evans, D. S., & Schmalensee, R. (2016). Matchmakers: The New Economics of Multisided Platforms. HBR Press.",
        "Fielding, R. T. (2000). Architectural Styles and the Design of Network-based Software Architectures. UC Irvine.",
        "Fowler, M. (2002). Patterns of Enterprise Application Architecture. Addison-Wesley.",
        "OWASP Foundation. (2021). OWASP Top Ten. https://owasp.org/www-project-top-ten/",
        "Pavlou, P. A., & Fygenson, M. (2006). MIS Quarterly, 30(1), 115–143.",
        "PostgreSQL Global Development Group. (2024). PostgreSQL 15 Documentation.",
        "Redis Ltd. (2024). Redis Documentation. https://redis.io/docs/",
        "Rochet, J.-C., & Tirole, J. (2003). Platform competition in two-sided markets. JEEA, 1(4), 990–1029.",
        "Marcotte, E. (2010). Responsive web design. A List Apart, 306.",
    ]
    for i, ref in enumerate(refs, 1):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.left_indent = Inches(0.25)
        p.paragraph_format.first_line_indent = Inches(-0.25)
        run = p.add_run(f"[{i}]  {ref}")
        run.font.size = Pt(10)
        run.font.name = "Times New Roman"


def build_article():
    print("Generating diagrams...")
    diagrams = generate_all_diagrams()

    doc = Document()
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1.1)
        section.right_margin = Inches(1.1)

    title_page(doc)
    setup_headers_footers(doc)
    abstract_section(doc)

    # ── 1. Introduction ──
    section_heading(doc, "1. Introduction", 1)
    body(doc,
         "The digital transformation of professional services has created substantial demand for platforms "
         "that connect clients with verified workers while reducing search friction, payment risk, and "
         "administrative overhead. Established marketplaces such as TaskRabbit, Thumbtack, and Upwork "
         "demonstrate that two-sided platforms can aggregate supply and demand at scale; however, local "
         "tradespeople often lack affordable channels to build reputation and receive secure online payments.")
    body(doc,
         "ProTech addresses this gap through an academically rigorous yet commercially informed implementation. "
         "Developed by Manye Patrice as the capstone for Web Technology in Business and Web Application "
         "Development, the platform serves clients who browse, book, and pay for services, and workers who "
         "manage profiles, accept jobs, and receive wallet-based payouts. The project evolved from a PHP/MySQL "
         "prototype into a Django-centred system aligned with enterprise web practice and the module rubric.")
    body(doc,
         "This article contributes four elements: (1) documentation of migration from PHP to domain-partitioned "
         "Django with a Jinja2 compatibility layer; (2) an internal escrow ledger design for marketplace trust; "
         "(3) explicit rubric-to-implementation mapping; and (4) a hybrid monetisation model grounded in "
         "two-sided market theory.")

    # ── 2. Related Work ──
    section_heading(doc, "2. Related Work", 1)
    body(doc,
         "Research on electronic commerce adoption emphasises perceived usefulness, ease of use, and trust "
         "as primary drivers of platform participation (Pavlou & Fygenson, 2006). Trust mechanisms—verified "
         "identity, reputation systems, and payment holding—reduce transaction risk (Edelman et al., 2012). "
         "ProTech operationalises these through role-separated dashboards, booking-linked reviews, and escrow "
         "holds released only upon job completion.")
    body(doc,
         "Platform economics literature formalises pricing on both sides of a market (Rochet & Tirole, 2003; "
         "Evans & Schmalensee, 2016). ProTech's design—free client access, worker commissions, optional "
         "subscriptions—follows asymmetric pricing common among services marketplaces.")
    body(doc,
         "Fielding (2000) and Fowler (2002) inform separation of presentation, domain logic, and persistence. "
         "Django's ORM, migrations, admin, and security middleware embody enterprise patterns. Docker Compose "
         "reflects reproducible DevOps practice. Security references OWASP Top Ten (2021): RBAC decorators, "
         "ORM queries, and PBKDF2 hashing address access control, injection, and cryptographic failures.")

    # ── 3. Architecture ──
    section_heading(doc, "3. System Architecture", 1)
    body(doc,
         "ProTech adopts a three-tier architecture deployed as four Docker Compose services: PostgreSQL (db), "
         "Redis (redis), Django/Gunicorn (backend), and Nginx (nginx). The presentation tier delivers "
         "server-rendered HTML via Jinja2 templates. The application tier partitions logic into users, bookings, "
         "reviews, and escrow applications. The data tier combines PostgreSQL 15 with Redis 7 for sessions "
         "and query caching. Django middleware enforces CSRF, authentication, and security headers on every request.")
    styled_table(doc, ["Tier", "Technology", "Responsibility"], [
        ("Presentation", "Jinja2 + Nginx", "UI rendering, static/media offload"),
        ("Application", "Django 5.2 + Gunicorn", "Business logic, auth, admin"),
        ("Persistence", "PostgreSQL 15", "Users, bookings, escrow ledger"),
        ("Cache / sessions", "Redis 7", "Worker search cache, sessions"),
        ("Orchestration", "Docker Compose", "Reproducible deployment"),
    ], col_widths=[1.3, 1.7, 3.5])
    caption(doc, "Table 1: Three-tier architecture and deployment components")

    # ── 4. Business Model ──
    section_heading(doc, "4. Business Model", 1)
    body(doc,
         "ProTech implements a hybrid model: transaction commissions at escrow release and tiered worker "
         "subscriptions for visibility. release_escrow() applies ESCROW_PLATFORM_FEE_PERCENT (default five "
         "percent). Tiered rates are specified for micro, standard, and commercial job bands. Subscription "
         "plans—Starter (free), Professional ($29/mo), Premium ($59/mo)—are documented but billing is not yet coded.")
    styled_table(doc, ["Plan", "Price", "Visibility", "Commission discount"], [
        ("Starter", "Free", "Basic listing; capped bookings", "None"),
        ("Professional", "$29/mo", "Priority above Starter", "−1% (4% effective)"),
        ("Premium", "$59/mo", "Featured placement", "−2% (3% effective)"),
    ], col_widths=[1.2, 1.0, 2.5, 1.8])
    caption(doc, "Table 2: Worker subscription tiers (specified; billing planned)")

    # ── 5. Security ──
    section_heading(doc, "5. Authentication and Security", 1)
    body(doc,
         "Authentication uses Django sessions with a custom EmailBackend. Passwords are PBKDF2-hashed. "
         "Registration creates client Users or User+Worker pairs. Authorisation uses user_type and "
         "@user_type_required to separate client, worker, and admin flows.")
    styled_table(doc, ["Threat (OWASP)", "Control", "Implementation"], [
        ("Broken access control", "RBAC decorator", "users/decorators.py"),
        ("Injection", "ORM parameterisation", "Django ORM"),
        ("Cryptographic failures", "PBKDF2 hashing", "Django auth"),
        ("CSRF", "CsrfViewMiddleware", "csrf_field() in templates"),
        ("XSS", "Auto-escaping", "Jinja2 templates"),
        ("Misconfiguration", "Environment secrets", ".env + .gitignore"),
    ], col_widths=[1.6, 1.5, 2.4])
    caption(doc, "Table 3: Security controls mapped to OWASP categories")

    section_heading(doc, "5.1 Contrast with Prior PHP Prototype", 2)
    body(doc,
         "The earlier PHP/MySQL prototype used bcrypt, MySQLi prepared statements, and PHP sessions in XAMPP. "
         "It lacked admin tooling, migrations, escrow, Docker deployment, automated tests, and caching. The "
         "Django rewrite consolidates these capabilities and replaces demo-mode URL bypasses with seed_demo "
         "and Docker entrypoint seeding.")

    doc.add_page_break()

    # ── 6. Escrow ──
    section_heading(doc, "6. Escrow Payment Subsystem", 1)
    body(doc,
         "ProTech holds client funds in escrow until work is complete. Logic resides in escrow/services.py. "
         "Workflow: deposit to UserWallet → create Booking → fund_escrow() creates EscrowHold → "
         "release_escrow() on completion (fee deducted) or refund_escrow() on cancellation.")
    add_figure(doc, diagrams["sequence"], width=5.9)
    caption(doc, "Figure 1: Escrow payment sequence diagram")
    code_box(doc, """users ──1:1── UserWallet
users ──1:1── Worker
users ──1:N── Booking ──1:1── EscrowHold
Booking ──1:1── Review
EscrowHold ── logs ── WalletTransaction
EscrowVault (singleton) — platform totals""")
    caption(doc, "Figure 2: Core entity relationships (textual ER summary)")
    add_figure(doc, diagrams["er"], width=5.9)
    caption(doc, "Figure 3: Entity-relationship diagram")

    # ── 7. Data Model ──
    section_heading(doc, "7. Data Model and Booking Workflow", 1)
    body(doc,
         "The schema centres on User, Worker, ServiceCategory, and Booking. Review attaches one-to-one to "
         "completed bookings. Message, Notification, Favorite, and WorkerAvailability support engagement.")
    styled_table(doc, ["Booking status", "Actor", "Escrow interaction"], [
        ("pending", "Worker accepts/rejects", "Client may fund escrow"),
        ("confirmed / in_progress", "Worker updates", "Funds held if funded"),
        ("completed", "Worker marks done", "release_escrow(); fee deducted"),
        ("cancelled", "Either party", "refund_escrow() if held"),
    ], col_widths=[1.5, 1.6, 2.4])
    caption(doc, "Table 4: Booking states and escrow coupling")

    # ── 8. DevOps ──
    section_heading(doc, "8. Implementation and DevOps", 1)
    body(doc,
         "Code is organised by domain: users/, bookings/, reviews/, escrow/, protech_project/. "
         "cache_utils.py caches worker search (300s) and categories (600s). Logging writes to stdout and "
         "logs/django.log. docker-entrypoint.sh runs migrate, collectstatic, seed_demo, and Gunicorn.")
    styled_table(doc, ["Rubric requirement", "Implementation", "Evidence"], [
        ("Django + templates", "Django 5.2 + Jinja2", "templates/, jinja2_env.py"),
        ("Auth & RBAC", "@user_type_required", "users/decorators.py"),
        ("Password reset", "Django token flow", "/password-reset/"),
        ("Admin / CMS", "Django Admin", "/admin/"),
        ("Docker + Redis", "Four-service Compose", "docker-compose.yml"),
        ("Tests + CI", "17 passing tests", "ci.yml"),
        ("Logging", "Per-app loggers", "settings.py"),
    ], col_widths=[1.5, 2.0, 2.8])
    caption(doc, "Table 5: Academic rubric alignment")

    # ── 9. Testing ──
    section_heading(doc, "9. Testing and Evaluation", 1)
    body(doc,
         "Unit tests validate models; integration tests cover login, RBAC, password reset, health endpoint, "
         "and escrow wallet debits. GitHub Actions runs flake8, migrations, full test suite, and /health/ check.")
    styled_table(doc, ["Metric", "Result", "Benchmark"], [
        ("Automated tests", "17 passing", "Module requirement met"),
        ("CI pipeline", "Lint, migrate, test, health", "GitHub Actions"),
        ("Deployment", "docker compose up --build", "Single-command demo"),
        ("Escrow commission", "5% default", "Live in services.py"),
        ("Cache", "Redis / locmem (CI)", "Environment-aware"),
    ], col_widths=[1.8, 2.0, 2.5])
    caption(doc, "Table 6: Evaluation metrics and outcomes")

    # ── 10–11. Limitations & Conclusion ──
    section_heading(doc, "10. Limitations", 1)
    body(doc,
         "ProTech uses demonstration wallet deposits, not a licensed payment provider. Subscriptions are "
         "specified but not billed. Messaging is not real-time. Browser E2E tests are absent. Growth "
         "projections are illustrative. Legacy Flask/PHP files remain for reference; Django is canonical.")

    section_heading(doc, "11. Conclusion and Future Work", 1)
    body(doc,
         "This paper presented ProTech as a production-oriented marketplace replacing a PHP prototype with "
         "Django, PostgreSQL, Redis, and Docker. The platform integrates discovery, booking, messaging, "
         "reviews, escrow, caching, logging, testing, and CI into a system suitable for master's assessment.")
    body(doc,
         "Future work prioritises Stripe integration, WorkerSubscription billing, WebSocket notifications, "
         "Selenium/Cypress E2E tests, geolocation matching, and two-factor authentication. The modular "
         "architecture supports these increments without fundamental restructuring.")

    doc.add_page_break()
    references_section(doc)

    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT)
    print(f"Article written to: {OUTPUT}")
    return OUTPUT


if __name__ == "__main__":
    build_article()
